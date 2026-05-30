"""docx_export.py — convert paper.md + figures to a styled DOCX file."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


def _strip_bold(text: str) -> str:
    """Remove ** markers and replace AI-tell dashes."""
    text = text.replace("**", "")
    text = text.replace("—", ", ")  # em dash → comma
    text = text.replace("–", "-")   # en dash → hyphen
    text = text.replace("―", "-")   # horizontal bar → hyphen
    return text


def export_to_docx(markdown_text: str, figures_dir: str, output_path: str) -> str:
    """Convert a markdown paper with embedded figures to DOCX.

    Args:
        markdown_text: The complete paper in Markdown format
        figures_dir: Directory containing PNG figures referenced in the markdown
        output_path: Where to save the .docx file

    Returns the output path on success.
    """
    figures_path = Path(figures_dir)
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Styles ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level in [1, 2, 3]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(14)
        elif level == 2:
            hs.font.size = Pt(12)
        else:
            hs.font.size = Pt(11)

    # ── Parse markdown ──────────────────────────────────────────────────
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Title line
        if line.startswith("# ") and not line.startswith("## "):
            title_text = _strip_bold(line[2:].strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = "Times New Roman"
            doc.add_paragraph()

        # Section headings — strip ** markers
        elif line.startswith("## "):
            doc.add_heading(_strip_bold(line[3:].strip()), level=1)

        elif line.startswith("### "):
            doc.add_heading(_strip_bold(line[4:].strip()), level=2)

        # Skip horizontal rules and empty lines
        elif line.strip() in ("---", "***", "___", "----", "- - -"):
            pass

        # Figures
        elif line.startswith("![") and "](" in line:
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                img_path = figures_path / m.group(2)
                if img_path.exists():
                    caption_text = ""
                    if i + 1 < len(lines) and lines[i + 1].startswith("*Figure"):
                        caption_text = lines[i + 1].strip("*").strip()
                        i += 1

                    if caption_text:
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cr = cap.add_run(caption_text)
                        cr.font.size = Pt(9)
                        cr.font.italic = True

                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        p_img.add_run().add_picture(str(img_path), width=Inches(5.0))
                    except Exception:
                        pass
                    doc.add_paragraph()

        elif line.startswith("*Figure"):
            pass  # Handled above

        # Lists
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            doc.add_paragraph(_strip_bold(line.strip()))

        elif re.match(r"^\d+\.\s", line.strip()):
            doc.add_paragraph(_strip_bold(line.strip()))

        # Regular paragraphs — strip **
        elif line.strip():
            doc.add_paragraph(_strip_bold(line))

        i += 1

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
